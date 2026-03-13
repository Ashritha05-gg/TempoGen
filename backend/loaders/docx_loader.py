# # loaders/docx_loader.py

# from docx import Document
# import os
# from typing import List, Dict

# def extract_docx_text(file_path: str) -> List[Dict]:
#     doc = Document(file_path)
#     output = []
#     filename = os.path.basename(file_path)

#     for i, para in enumerate(doc.paragraphs):
#         text = para.text.strip()
#         if text:
#             output.append({
#                 "text": text,
#                 "page": f"paragraph-{i+1}",
#                 "source": filename
#             })

#     return output

from docx import Document


def extract_docx_text(file_path: str):
    """
    Extract text from DOCX including tables.
    Returns: list of { 'text': str, 'page': int }
    """

    doc = Document(file_path)

    pages = []
    page_num = 1

    full_text = []

    # ---------- Paragraphs ----------
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            full_text.append(text)

    # ---------- Tables ----------
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                full_text.append(" | ".join(row_text))

    if not full_text:
        return []

    pages.append({
        "text": "\n".join(full_text),
        "page": page_num
    })

    return pages
